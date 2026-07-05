from pathlib import Path


def extract_text(file_path: Path | str, filename: str = "") -> str:
    """Extract text from PDF, DOCX, or plain text files."""
    file_path = Path(file_path)
    suffix = file_path.suffix.lower() if file_path.suffix else Path(filename).suffix.lower()

    if suffix == ".pdf":
        text = _extract_pdf_pymupdf(file_path)
        if not text.strip():
            text = _extract_pdf_pdfplumber(file_path)
        if not text.strip():
            raise ValueError(
                "无法从 PDF 中提取文本。文件可能是扫描件（图片）或使用了特殊字体。"
                "请尝试直接粘贴简历内容，或使用文字版 PDF。"
            )
        return text
    elif suffix in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif suffix in (".txt", ".md", ".text"):
        return file_path.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"不支持的文件格式: {suffix}，请上传 PDF、DOCX 或 TXT 文件。")


def _extract_pdf_pymupdf(path: Path) -> str:
    """Primary PDF parser using pymupdf (handles Chinese/unicode well)."""
    import fitz  # pymupdf

    text_parts: list[str] = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            text = page.get_text()
            if text.strip():
                text_parts.append(text)
    return "\n\n".join(text_parts)


def _extract_pdf_pdfplumber(path: Path) -> str:
    """Fallback PDF parser using pdfplumber (suppresses font warnings)."""
    import logging
    import pdfplumber

    # Suppress pdfminer font warnings
    logging.getLogger("pdfminer").setLevel(logging.ERROR)
    logging.getLogger("pdfplumber").setLevel(logging.ERROR)

    text_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _extract_docx(path: Path) -> str:
    """Extract text from DOCX files."""
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        # Also try tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
    if not paragraphs:
        raise ValueError("文档中没有找到文本内容。")
    return "\n".join(paragraphs)
